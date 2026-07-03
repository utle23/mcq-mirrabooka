"""Packaging Order — supplier-aware order composer.

Each supplier has a fixed item catalogue (codes, names, units) and contact
details. The main page lets the staff member tick / set quantities, then the
single "Review & Send Order" action shows the exact Word document content and,
on confirmation, emails the supplier through Brevo with the .docx auto-attached
(a hidden copy is BCC'd to the owner).

Admin can edit supplier contact details, add / edit / delete catalogue
items, and tweak default quantities directly from the same screen.
"""
from __future__ import annotations

import base64
import io
import json
import os
import re
import sqlite3
from datetime import datetime, date, timedelta
from functools import wraps

from flask import (Blueprint, render_template, request, redirect, url_for,
                   session, jsonify, send_file)

from store_scope import current_store_id, store_filter_clause

packaging_bp = Blueprint('packaging', __name__, url_prefix='/packaging')
DB_PATH: str | None = None
SUBIACO_STORE_ID = 3
SUBIACO_PACKAGING_DELIVERY_DAYS = 'TUE,THU'


# ── Catalogue (sourced from the Word order template) ─────────────────────────

JACCUS_ITEMS_SEED = [
    # (product_code, name_en, name_vi, unit, default_qty)
    ('NapQBES',             'Quilted Brown Express (Tork Xpress Dispenser) Napkin', 'Khăn giấy dispenser Tork Xpress', 'carton', 0),
    ('1WFB',                'Tui banh ngot',                              'Túi bánh ngọt',                                    'bag',    4),
    ('KCB-M',               'Hop nho — catering box',                     'Hộp nhỏ — catering box',                           'carton', 0),
    ('KCBWLid-M',           'Nap hop nho — catering box',                 'Nắp hộp nhỏ — catering box',                       'carton', 0),
    ('KCB-L',               'Hop lon — catering box',                     'Hộp lớn — catering box',                           'carton', 0),
    ('KCBWLid-L',           'Nap hop lon — catering box',                 'Nắp hộp lớn — catering box',                       'carton', 0),
    ('KDTR-4-PLA',          'Hop goi cuon',                               'Hộp gỏi cuốn',                                     'carton', 0),
    ('KDTR-4Lid',           'Nap hop goi cuon',                           'Nắp hộp gỏi cuốn',                                 'carton', 0),
    ('P200',                'Hop nuoc mam',                               'Hộp nước mắm',                                     'carton', 0),
    ('P200Lid',             'Nap hop nuoc mam',                           'Nắp hộp nước mắm',                                 'carton', 0),
    ('Rec1000-PLA-K',       'Hop vuong take away',                        'Hộp vuông take away',                              'carton', 0),
    ('RecPaper-PPLid',      'Nap hop vuong take away',                    'Nắp hộp vuông take away',                          'carton', 0),
    ('Rd24-PLA-W',          'Hop soup trang',                             'Hộp soup trắng',                                   'carton', 0),
    ('RdPPLid-115-F',       'Nap hop soup trang',                         'Nắp hộp soup trắng',                               'carton', 0),
    ('PaperBowl-Extra Large','Hop bun heo quay',                          'Hộp bún heo quay',                                 'carton', 2),
    ('BPB-PETLid184',       'Nap hop bun heo quay',                       'Nắp hộp bún heo quay',                             'carton', 1),
    ('EC-DCC390',           'Ly cafe',                                    'Ly cà phê',                                        'carton', 0),
    ('EC-DCC500',           'Ly juice',                                   'Ly juice',                                         'carton', 0),
    ('BioBCL-90C-Pulp-F',   'Nap ly',                                     'Nắp ly',                                           'carton', 0),
    ('DSPaperBlk',          'Ong hut',                                    'Ống hút',                                          'carton', 0),
    ('ChopstickBam',        'Dua',                                        'Đũa',                                              'box',    0),
    ('WoodenFrk',           'Nia',                                        'Nĩa',                                              'box',    1),
    ('WoodenKnf',           'Dao',                                        'Dao',                                              'box',    0),
    ('WoodenSpn',           'Muong',                                      'Muỗng',                                            'box',    0),
    ('PulpCSpn',            'Muong soup',                                 'Muỗng soup',                                       'box',    1),
    ('NitrileBluPF-Md',     'Glove M — bao tay',                          'Găng tay M',                                       'carton', 2),
    ('NitrileBluPF-Lg',     'Glove L — bao tay',                          'Găng tay L',                                       'carton', 2),
    ('BL82/35',             'Bao rac den',                                'Bao rác đen',                                      'box',    0),
    ('HTSlimline',          'Khan giay lau tay',                          'Khăn giấy lau tay',                                'carton', 0),
    ('CUP-HOLDER',          'Cup holder',                                 'Khay đựng ly',                                     'carton', 0),
    ('Surplus20',           'Nuoc rua chen 20L',                          'Nước rửa chén 20L',                                'can',    0),
    ('Oven5',               'Nuoc rua lo nuong 5L',                       'Nước rửa lò nướng 5L',                             'can',    0),
    ('AF44/150',            'Foil — giay bac',                            'Giấy bạc',                                         'each',   0),
    ('CW45/600 Pro',        'Clingwrap 45cm x 600m',                      'Màng bọc thực phẩm 45cm x 600m',                   'each',   0),
    ('BioR-500Y',           '500ml Clear Biocup',                         'Cốc Biocup 500ml trong',                           'carton', 0),
    ('BioC-96D(N)',         'Dome Lid (no hole) for 300-700ml BioCup',    'Nắp vòm BioCup 300-700ml',                         'carton', 0),
    ('LWGP33x40',           'Wrap banh mi',                               'Túi gói bánh mì',                                  'bag',    0),
    ('Blitz5',              'Blitz Multi-Purpose Floor Cleaner & Degreaser','Blitz — chất tẩy đa năng',                       'can',    0),
]

JACCUS_SUPPLIER_SEED = {
    'name':           'Jaccus Trading',
    'email':          'Orders@jaccus.com.au',
    'phone':          '08-9248 9668',
    'cc_emails':      'Mirrabooka@mcqinternational.com',
    'delivery_days':  'WED,FRI',
    'cafe_name':      'MCQ Vietnamese Street Food — MIRRABOOKA',
    'cafe_address':   'Shop MM4/43 Yirrigan Dr, Mirrabooka WA 6061',
    'cafe_contacts':  'Khoi: 0449819235',
    'notes':          '',
}

PACKAGING_ITEM_UNIT_FIXES = {
    'AF44/150': 'each',
    'CW45/600 Pro': 'each',
    'BioR-500Y': 'carton',
    'BioC-96D(N)': 'carton',
}

JACCUS_PRICE_DATA = {
    # product_code: (unit_of_measure from Jaccus price list, price per order unit)
    'NapQBES': ('6000', 49.90),
    '1WFB': ('1000', 19.90),
    'KCB-M': ('100', 60.90),
    'KCBWLid-M': ('100', 49.00),
    'KCB-L': ('50', 44.90),
    'KCBWLid-L': ('50', 37.00),
    'KDTR-4-PLA': ('400', 69.90),
    'KDTR-4Lid': ('400', 53.00),
    'P200': ('3000', 87.90),
    'P200Lid': ('3000', 74.00),
    'Rec1000-PLA-K': ('300', 51.90),
    'RecPaper-PPLid': ('300', 25.00),
    'Rd24-PLA-W': ('500', 76.90),
    'RdPPLid-115-F': ('500', 31.00),
    'PaperBowl-Extra Large': ('300', 78.90),
    'BPB-PETLid184': ('300', 47.00),
    'EC-DCC390': ('1000', 109.90),
    'EC-DCC500': ('1000', 124.90),
    'BioBCL-90C-Pulp-F': ('1000', 66.00),
    'DSPaperBlk': ('2500', 36.90),
    'ChopstickBam': ('3000', 49.90),
    'WoodenFrk': ('1000', 21.90),
    'WoodenKnf': ('1000', 18.90),
    'WoodenSpn': ('1000', 22.90),
    'PulpCSpn': ('1000', 49.90),
    'NitrileBluPF-Md': ('10 Box', 69.90),
    'NitrileBluPF-Lg': ('10 Box', 69.90),
    'BL82/35': ('200', 38.90),
    'HTSlimline': ('4000', 44.90),
    'Surplus20': ('20lt', 33.90),
    'Oven5': ('5lt', 29.90),
    'AF44/150': ('Each', 22.90),
    'CW45/600 Pro': ('1 Roll', 23.90),
    'BioR-500Y': ('1000', 219.90),
    'BioC-96D(N)': ('1000', 94.90),
    'LWGP33x40': ('800', 19.90),
    'Blitz5': ('5lt', 22.90),
    'PB#16': ('250', 36.90),
    'WipesBlu (Ea)': ('1 Roll', 11.90),
    'BL120/35': ('100', 23.90),
    'BL120/35 (100)': ('100', 23.90),
    'BL240/35': ('100', 54.90),
    'RT80': ('16 Roll', 49.90),
    'Bleach5': ('5lt', 12.90),
    'JetDryPlus15': ('15lt', 79.90),
    'JetKlean20': ('20lt', 89.90),
    'Sanitiser5': ('5lt', 26.90),
    'Surplus5': ('5lt', 16.90),
    'Bake40/120 (Ea)': ('Each', 26.90),
    '4CupPulp (300)': ('300', 60.90),
}

DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

# Every packaging order is silently copied here so the owner keeps a full record
# of what each branch sends to the supplier (hidden from the supplier via BCC).
OWNER_BCC_EMAIL = 'utle23.23@gmail.com'

# MCQ Vietnamese Street Food logo, embedded at the top of the Word order.
LOGO_PATH = os.path.join(os.path.dirname(__file__), 'static', 'logo.png')

DAY_CODES = {
    'MON': 0, 'MONDAY': 0,
    'TUE': 1, 'TUES': 1, 'TUESDAY': 1,
    'WED': 2, 'WEDNESDAY': 2,
    'THU': 3, 'THUR': 3, 'THURS': 3, 'THURSDAY': 3,
    'FRI': 4, 'FRIDAY': 4,
    'SAT': 5, 'SATURDAY': 5,
    'SUN': 6, 'SUNDAY': 6,
}


# ── Helpers ──────────────────────────────────────────────────────────────────

def _conn():
    c = sqlite3.connect(DB_PATH, timeout=30)
    c.row_factory = sqlite3.Row
    return c


def _money(amount) -> str:
    try:
        return f"${float(amount or 0):,.2f}"
    except (TypeError, ValueError):
        return "$0.00"


def _float_form(value, default=0.0):
    try:
        return max(0.0, float(value or default))
    except (TypeError, ValueError):
        return default


def _store_delivery_days_for_supplier(supplier: dict) -> str:
    """Return the effective supplier delivery days after branch-level rules."""
    try:
        store_id = int(supplier.get('store_id') or current_store_id())
    except Exception:
        store_id = current_store_id()
    if store_id == SUBIACO_STORE_ID:
        return SUBIACO_PACKAGING_DELIVERY_DAYS
    return (supplier.get('delivery_days') or '').strip().upper()


def _with_effective_delivery_days(supplier: dict) -> dict:
    supplier = dict(supplier)
    supplier['delivery_days'] = _store_delivery_days_for_supplier(supplier)
    return supplier


def _normalised_delivery_days(raw: str) -> list[tuple[str, int]]:
    days: list[tuple[str, int]] = []
    seen: set[int] = set()
    for token in re.split(r'[,/; ]+', (raw or '').upper()):
        if not token:
            continue
        weekday = DAY_CODES.get(token)
        if weekday is None or weekday in seen:
            continue
        seen.add(weekday)
        code = next(k for k, v in DAY_CODES.items() if v == weekday and len(k) == 3)
        days.append((code, weekday))
    return days


def _next_delivery_date(delivery_days: str, today_value: date | None = None) -> str:
    """Next configured delivery day after today, never same-day."""
    today_dt = today_value or date.today()
    best: date | None = None
    for _, weekday in _normalised_delivery_days(delivery_days):
        delta = (weekday - today_dt.weekday()) % 7
        if delta == 0:
            delta = 7
        candidate = today_dt + timedelta(days=delta)
        if best is None or candidate < best:
            best = candidate
    return (best or today_dt).isoformat()


def _resolve_delivery_date(supplier: dict, requested: str) -> str:
    delivery_days = _store_delivery_days_for_supplier(supplier)
    requested = (requested or '').strip()
    allowed = _normalised_delivery_days(delivery_days)
    try:
        requested_date = datetime.strptime(requested, '%Y-%m-%d').date()
    except (TypeError, ValueError):
        requested_date = None

    if requested_date:
        if not allowed:
            return requested_date.isoformat()
        today_dt = date.today()
        allowed_weekdays = {weekday for _, weekday in allowed}
        if requested_date > today_dt and requested_date.weekday() in allowed_weekdays:
            return requested_date.isoformat()

    return _next_delivery_date(delivery_days)


def _sync_jaccus_catalog_prices(c):
    # Keep the original Mirrabooka seed catalogue complete, but apply price data
    # to every branch's Jaccus catalogue by product code.
    supplier = c.execute(
        "SELECT id FROM packaging_suppliers WHERE name=? AND active=1 AND store_id=1 ORDER BY id LIMIT 1",
        (JACCUS_SUPPLIER_SEED['name'],)
    ).fetchone()
    if supplier:
        sid = supplier['id']
        existing_codes = {
            r['product_code']
            for r in c.execute(
                "SELECT product_code FROM packaging_items WHERE supplier_id=? AND active=1",
                (sid,)
            ).fetchall()
        }
        next_sort = c.execute(
            "SELECT COALESCE(MAX(sort_order), -1) + 1 AS n FROM packaging_items WHERE supplier_id=?",
            (sid,)
        ).fetchone()['n']

        for code, en, vi, unit, qty in JACCUS_ITEMS_SEED:
            if code in existing_codes:
                continue
            c.execute(
                '''INSERT INTO packaging_items
                   (supplier_id, product_code, name_en, name_vi, unit, default_qty, sort_order)
                   VALUES (?,?,?,?,?,?,?)''',
                (sid, code, en, vi, unit, qty, next_sort)
            )
            existing_codes.add(code)
            next_sort += 1

    for supplier_row in c.execute(
        "SELECT id FROM packaging_suppliers WHERE name=? AND active=1",
        (JACCUS_SUPPLIER_SEED['name'],)
    ).fetchall():
        sid = supplier_row['id']
        for product_code, (unit_measure, unit_price) in JACCUS_PRICE_DATA.items():
            c.execute(
                '''UPDATE packaging_items
                   SET unit_measure=?, unit_price=?
                   WHERE supplier_id=? AND product_code=? AND active=1''',
                (unit_measure, unit_price, sid, product_code)
            )


def _login_required(f):
    @wraps(f)
    def d(*a, **kw):
        if not session.get('logged_in'):
            return redirect(url_for('login_page'))
        return f(*a, **kw)
    return d


def _admin_required(f):
    @wraps(f)
    def d(*a, **kw):
        if not session.get('logged_in'):
            return redirect(url_for('login_page'))
        if session.get('role') not in ('admin', 'super_admin'):
            return render_template('access_denied.html'), 403
        return f(*a, **kw)
    return d


def init_packaging(db_path: str):
    global DB_PATH
    DB_PATH = db_path
    with _conn() as c:
        c.executescript('''
            CREATE TABLE IF NOT EXISTS packaging_suppliers (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                name          TEXT NOT NULL,
                email         TEXT NOT NULL DEFAULT '',
                phone         TEXT NOT NULL DEFAULT '',
                cc_emails     TEXT NOT NULL DEFAULT '',
                delivery_days TEXT NOT NULL DEFAULT '',
                cafe_name     TEXT NOT NULL DEFAULT 'MCQ Vietnamese Street Food — MIRRABOOKA',
                cafe_address  TEXT NOT NULL DEFAULT '',
                cafe_contacts TEXT NOT NULL DEFAULT '',
                notes         TEXT NOT NULL DEFAULT '',
                active        INTEGER NOT NULL DEFAULT 1,
                sort_order    INTEGER NOT NULL DEFAULT 0,
                store_id      INTEGER NOT NULL DEFAULT 1
            );
            CREATE TABLE IF NOT EXISTS packaging_items (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                supplier_id   INTEGER NOT NULL REFERENCES packaging_suppliers(id) ON DELETE CASCADE,
                product_code  TEXT NOT NULL DEFAULT '',
                name_en       TEXT NOT NULL,
                name_vi       TEXT NOT NULL DEFAULT '',
                unit          TEXT NOT NULL DEFAULT 'carton',
                unit_measure  TEXT NOT NULL DEFAULT '',
                unit_price    REAL NOT NULL DEFAULT 0,
                default_qty   INTEGER NOT NULL DEFAULT 0,
                sort_order    INTEGER NOT NULL DEFAULT 0,
                active        INTEGER NOT NULL DEFAULT 1
            );
            CREATE TABLE IF NOT EXISTS packaging_orders (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                supplier_id   INTEGER NOT NULL REFERENCES packaging_suppliers(id),
                delivery_date TEXT NOT NULL DEFAULT '',
                composed_by   TEXT NOT NULL DEFAULT '',
                sent_at       TEXT DEFAULT (datetime('now','localtime')),
                send_channel  TEXT NOT NULL DEFAULT '',   -- gmail / brevo / mailto / preview
                subject       TEXT NOT NULL DEFAULT '',
                body          TEXT NOT NULL DEFAULT '',
                payload_json  TEXT NOT NULL DEFAULT ''    -- chosen items snapshot
            );
        ''')

        # Migration: add cafe_address to pre-existing supplier tables.
        cols = [r['name'] for r in c.execute("PRAGMA table_info(packaging_suppliers)").fetchall()]
        if 'cafe_address' not in cols:
            c.execute("ALTER TABLE packaging_suppliers ADD COLUMN cafe_address TEXT NOT NULL DEFAULT ''")
        if 'store_id' not in cols:
            c.execute("ALTER TABLE packaging_suppliers ADD COLUMN store_id INTEGER NOT NULL DEFAULT 1")
        item_cols = [r['name'] for r in c.execute("PRAGMA table_info(packaging_items)").fetchall()]
        if 'unit_measure' not in item_cols:
            c.execute("ALTER TABLE packaging_items ADD COLUMN unit_measure TEXT NOT NULL DEFAULT ''")
        if 'unit_price' not in item_cols:
            c.execute("ALTER TABLE packaging_items ADD COLUMN unit_price REAL NOT NULL DEFAULT 0")
        # Backfill the MCQ Mirrabooka delivery address where it's still blank.
        c.execute("UPDATE packaging_suppliers SET cafe_address=? WHERE COALESCE(cafe_address,'')=''",
                  (JACCUS_SUPPLIER_SEED['cafe_address'],))
        # Subiaco packaging deliveries are fixed to Tuesday/Thursday. Keep this
        # branch rule in the DB too so old rows do not render as "today".
        c.execute(
            "UPDATE packaging_suppliers SET delivery_days=? WHERE store_id=? AND active=1",
            (SUBIACO_PACKAGING_DELIVERY_DAYS, SUBIACO_STORE_ID)
        )
        # Only seed the Mirrabooka (store_id=1) contact. Other branches keep
        # their own contact (e.g. Subiaco's Kenny Ho) instead of being clobbered
        # with Mirrabooka's on every startup.
        c.execute(
            "UPDATE packaging_suppliers SET cafe_contacts=? "
            "WHERE name=? AND active=1 AND store_id=1",
            (JACCUS_SUPPLIER_SEED['cafe_contacts'], JACCUS_SUPPLIER_SEED['name'])
        )
        for product_code, unit in PACKAGING_ITEM_UNIT_FIXES.items():
            c.execute(
                "UPDATE packaging_items SET unit=? WHERE product_code=? AND active=1",
                (unit, product_code)
            )

        # Seed Jaccus if no suppliers exist yet
        if c.execute('SELECT COUNT(*) c FROM packaging_suppliers').fetchone()['c'] == 0:
            cur = c.execute('''INSERT INTO packaging_suppliers
                (name, email, phone, cc_emails, delivery_days, cafe_name, cafe_address, cafe_contacts)
                VALUES (?,?,?,?,?,?,?,?)''',
                (JACCUS_SUPPLIER_SEED['name'],
                 JACCUS_SUPPLIER_SEED['email'],
                 JACCUS_SUPPLIER_SEED['phone'],
                 JACCUS_SUPPLIER_SEED['cc_emails'],
                 JACCUS_SUPPLIER_SEED['delivery_days'],
                 JACCUS_SUPPLIER_SEED['cafe_name'],
                 JACCUS_SUPPLIER_SEED['cafe_address'],
                 JACCUS_SUPPLIER_SEED['cafe_contacts']))
            sid = cur.lastrowid
            for i, (code, en, vi, unit, qty) in enumerate(JACCUS_ITEMS_SEED):
                c.execute('''INSERT INTO packaging_items
                    (supplier_id, product_code, name_en, name_vi, unit, default_qty, sort_order)
                    VALUES (?,?,?,?,?,?,?)''',
                    (sid, code, en, vi, unit, qty, i))
        _sync_jaccus_catalog_prices(c)


# ── Data helpers ─────────────────────────────────────────────────────────────

def _suppliers():
    with _conn() as c:
        rows = c.execute(
            'SELECT * FROM packaging_suppliers WHERE active=1 AND store_id=? '
            'ORDER BY sort_order, name', (current_store_id(),)).fetchall()
    return [_with_effective_delivery_days(dict(r)) for r in rows]


def _supplier(sid):
    with _conn() as c:
        row = c.execute('SELECT * FROM packaging_suppliers WHERE id=? AND store_id=?',
                        (sid, current_store_id())).fetchone()
    return _with_effective_delivery_days(dict(row)) if row else None


def _items(sid):
    with _conn() as c:
        rows = c.execute(
            'SELECT * FROM packaging_items WHERE supplier_id=? AND active=1 '
            'ORDER BY sort_order, id', (sid,)).fetchall()
    return [dict(r) for r in rows]


# ── Email composition ────────────────────────────────────────────────────────

def _format_delivery_date(s: str) -> str:
    try:
        return datetime.strptime(s, '%Y-%m-%d').strftime('%A, %d/%m/%Y')
    except Exception:
        return s or '-'


def _brand_name(supplier: dict) -> str:
    return (supplier.get('cafe_name') or 'MCQ Restaurant').strip()


def _delivery_line(supplier: dict) -> str:
    brand = _brand_name(supplier)
    address = (supplier.get('cafe_address') or '').strip()
    if not address:
        return brand
    brand_root = re.split(r'\s+[—-]\s+', brand, maxsplit=1)[0].strip().lower()
    if address.lower().startswith(brand.lower()) or (brand_root and address.lower().startswith(brand_root)):
        return address
    return f"{brand} - {address}"


def _order_subject(supplier: dict, delivery_date: str) -> str:
    deliv_pretty = _format_delivery_date(delivery_date)
    return f"{_brand_name(supplier)} - Packaging order - delivery {deliv_pretty}"


def _order_filename(supplier: dict, delivery_date: str) -> str:
    brand = re.sub(r'[^A-Za-z0-9]+', '-', _brand_name(supplier)).strip('-') or 'MCQ'
    date_part = re.sub(r'[^0-9A-Za-z-]+', '-', delivery_date or date.today().isoformat()).strip('-')
    return f"{brand}-packaging-order-{date_part}.docx"


def _order_totals(items_with_qty: list[dict]) -> tuple[int, float]:
    total_qty = sum(i['qty'] for i in items_with_qty)
    total_money = sum(i['qty'] * float(i.get('unit_price') or 0) for i in items_with_qty)
    return total_qty, total_money


def _selected_order_items(sid: int) -> list[dict]:
    chosen = []
    for it in _items(sid):
        raw = request.form.get(f'qty_{it["id"]}', '').strip()
        try:
            qty = int(raw or 0)
        except ValueError:
            qty = 0
        if qty > 0:
            chosen.append({**it, 'qty': qty})
    return chosen


def _compose_order(supplier: dict, items_with_qty: list[dict],
                    delivery_date: str, extra_note: str = '') -> tuple[str, str]:
    """Return a short, professional email body.

    The official order details live in the attached Word document, so the email
    body stays minimal — greeting, the attachment/delivery line, and where to
    deliver. Every value is pulled from the supplier record, so editing the cafe
    name/address updates both this email and the Word attachment.
    """
    deliv_pretty = _format_delivery_date(delivery_date)
    subject = _order_subject(supplier, delivery_date)

    cafe_address = (supplier.get('cafe_address') or '').strip()
    parts = [
        f"Hello {supplier['name']},",
        "",
        f"Please see the attached Microsoft Word packaging order for delivery on {deliv_pretty}.",
    ]
    if cafe_address:
        parts += ["", f"Deliver to: {_delivery_line(supplier)}"]

    if extra_note.strip():
        parts += ["", "Note:", extra_note.strip()]

    return subject, '\n'.join(parts)


def _build_order_docx(supplier: dict, items_with_qty: list[dict],
                      delivery_date: str, extra_note: str = '',
                      composed_by: str = '') -> bytes:
    """Build a professional Word order attachment from the selected quantities."""
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    GREEN = RGBColor(0x3F, 0x8F, 0x3A)
    DARK = RGBColor(0x2B, 0x2B, 0x2B)
    MUTED = RGBColor(0x77, 0x82, 0x77)
    ACCENT_HEX = '3F8F3A'
    HEADER_FILL = 'E9F3E6'
    STRIPE_FILL = 'F6F9F5'
    LABEL_FILL = 'F2F4F6'
    BORDER_HEX = 'D9E2D6'

    def set_font(run, size=10.5, bold=False, color=None):
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        run.font.size = Pt(size)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color

    def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in('w:tcMar')
        if tc_mar is None:
            tc_mar = OxmlElement('w:tcMar')
            tc_pr.append(tc_mar)
        for m, v in (('top', top), ('bottom', bottom), ('start', start), ('end', end)):
            node = tc_mar.find(qn(f'w:{m}'))
            if node is None:
                node = OxmlElement(f'w:{m}')
                tc_mar.append(node)
            node.set(qn('w:w'), str(v))
            node.set(qn('w:type'), 'dxa')

    def shade(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tc_pr.append(shd)
        shd.set(qn('w:fill'), fill)

    def set_cell(cell, text, *, bold=False, size=9.2, color=None,
                 align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
        cell.text = ''
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        if fill:
            shade(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.10
        run = p.add_run(str(text or ''))
        set_font(run, size=size, bold=bold, color=color)

    def set_table_width(table, widths):
        table.autofit = False
        dxa_widths = [int(round(width * 1440)) for width in widths]
        for row in table.rows:
            for idx, dxa in enumerate(dxa_widths):
                cell = row.cells[idx]
                cell.width = Inches(widths[idx])
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn('w:tcW'))
                if tc_w is None:
                    tc_w = OxmlElement('w:tcW')
                    tc_pr.append(tc_w)
                tc_w.set(qn('w:w'), str(dxa))
                tc_w.set(qn('w:type'), 'dxa')
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is None:
            tbl_w = OxmlElement('w:tblW')
            tbl_pr.append(tbl_w)
        tbl_w.set(qn('w:w'), '9360')
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_ind = tbl_pr.find(qn('w:tblInd'))
        if tbl_ind is None:
            tbl_ind = OxmlElement('w:tblInd')
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn('w:w'), '120')
        tbl_ind.set(qn('w:type'), 'dxa')
        tbl_grid = table._tbl.tblGrid
        for child in list(tbl_grid):
            tbl_grid.remove(child)
        for dxa in dxa_widths:
            col = OxmlElement('w:gridCol')
            col.set(qn('w:w'), str(dxa))
            tbl_grid.append(col)

    def set_table_borders(table, *, color=BORDER_HEX, sz=4, val='single'):
        tbl_pr = table._tbl.tblPr
        existing = tbl_pr.find(qn('w:tblBorders'))
        if existing is not None:
            tbl_pr.remove(existing)
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'), val)
            e.set(qn('w:sz'), str(sz))
            e.set(qn('w:space'), '0')
            e.set(qn('w:color'), color)
            borders.append(e)
        tbl_pr.append(borders)

    def bottom_rule(paragraph, *, color=ACCENT_HEX, sz=12):
        p_pr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(sz))
        bottom.set(qn('w:space'), '6')
        bottom.set(qn('w:color'), color)
        pbdr.append(bottom)
        p_pr.append(pbdr)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    # Footer — single, quiet line.
    footer = section.footer.paragraphs[0]
    footer.text = ''
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"{_brand_name(supplier)}  ·  Packaging order generated by the MCQ Food Safety App")
    set_font(fr, size=8, color=MUTED)

    # ── Letterhead: logo + brand name, with an accent divider underneath ──
    banner = doc.add_table(rows=1, cols=2)
    set_table_width(banner, [1.7, 4.8])
    set_table_borders(banner, val='nil')
    logo_cell = banner.rows[0].cells[0]
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(logo_cell, top=0, bottom=0, start=0, end=0)
    lp = logo_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        lp.add_run().add_picture(LOGO_PATH, width=Inches(1.4))
    except Exception:
        set_font(lp.add_run('MCQ'), size=20, bold=True, color=GREEN)

    info_cell = banner.rows[0].cells[1]
    info_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(info_cell)
    bp = info_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    bp.paragraph_format.space_after = Pt(0)
    set_font(bp.add_run(_brand_name(supplier)), size=14, bold=True, color=GREEN)
    tag = info_cell.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tag.paragraph_format.space_before = Pt(0)
    set_font(tag.add_run('Vietnamese Street Food'), size=9, color=MUTED)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_before = Pt(2)
    divider.paragraph_format.space_after = Pt(10)
    bottom_rule(divider)

    # ── Title + delivery date ──
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    set_font(title.add_run('PACKAGING ORDER'), size=20, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    set_font(subtitle.add_run(f"Delivery: {_format_delivery_date(delivery_date)}"),
             size=11, bold=True, color=GREEN)

    # ── Order details (clean, borderless 2-column block) ──
    meta_rows = [
        ('Supplier', supplier.get('name') or ''),
        ('Deliver to', _delivery_line(supplier)),
        ('Contact', (supplier.get('cafe_contacts') or '').replace('\r\n', '\n')),
        ('Prepared by', composed_by or '-'),
        ('Order date', datetime.now().strftime('%d/%m/%Y')),
    ]
    meta = doc.add_table(rows=len(meta_rows), cols=2)
    set_table_width(meta, [1.35, 5.15])
    set_table_borders(meta, val='nil')
    for row, (label, value) in zip(meta.rows, meta_rows):
        set_cell(row.cells[0], label, bold=True, size=9.5, color=GREEN, fill=LABEL_FILL)
        set_cell(row.cells[1], value, size=9.5, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Line items ──
    table = doc.add_table(rows=1, cols=7)
    widths = [1.05, 2.0, 0.55, 0.68, 0.72, 0.75, 0.75]
    set_table_width(table, widths)
    set_table_borders(table)
    headings = ['Product Code', 'Name (Vietnamese)', 'Qty', 'Unit', 'Measure', 'Unit Price', 'Line Total']
    aligns = [
        WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
    ]
    for cell, text, align in zip(table.rows[0].cells, headings, aligns):
        set_cell(cell, text, bold=True, size=8.8, color=DARK, align=align, fill=HEADER_FILL)

    total_money = 0.0
    for idx, it in enumerate(items_with_qty):
        qty = int(it.get('qty') or 0)
        unit_price = float(it.get('unit_price') or 0)
        line_total = qty * unit_price
        total_money += line_total
        stripe = STRIPE_FILL if idx % 2 else None
        values = [
            it.get('product_code') or '',
            it.get('name_vi') or it.get('name_en') or '',
            str(qty),
            it.get('unit') or '',
            it.get('unit_measure') or '',
            _money(unit_price),
            _money(line_total),
        ]
        row = table.add_row().cells
        for cell, text, align in zip(row, values, aligns):
            set_cell(cell, text, size=9, color=DARK, align=align, fill=stripe)

    # Total row inside the table so it lines up under Line Total.
    total_row = table.add_row().cells
    set_cell(total_row[0], 'TOTAL ORDER VALUE', bold=True, size=9.5, color=DARK, fill=HEADER_FILL)
    for c in total_row[1:6]:
        set_cell(c, '', fill=HEADER_FILL)
    set_cell(total_row[6], _money(total_money), bold=True, size=10, color=GREEN,
             align=WD_ALIGN_PARAGRAPH.RIGHT, fill=HEADER_FILL)

    if extra_note.strip():
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(10)
        set_font(note.add_run('Note: '), size=9.5, bold=True, color=GREEN)
        set_font(note.add_run(extra_note.strip()), size=9.5, color=DARK)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Routes ───────────────────────────────────────────────────────────────────

@packaging_bp.route('/')
@_login_required
def packaging_home():
    suppliers = _suppliers()
    if not suppliers:
        return render_template('packaging.html', supplier=None, suppliers=[], items=[])

    sid_str = request.args.get('supplier', '')
    try:
        sid = int(sid_str) if sid_str else suppliers[0]['id']
    except ValueError:
        sid = suppliers[0]['id']
    supplier = _supplier(sid) or suppliers[0]
    items    = _items(supplier['id'])

    # Pick a sensible default delivery date: the next configured delivery day.
    deliv_default = _next_delivery_date(supplier.get('delivery_days') or '')

    return render_template('packaging.html',
        supplier=supplier, suppliers=suppliers, items=items,
        delivery_default=deliv_default, is_admin=session.get('role') in ('admin','super_admin'))


@packaging_bp.route('/compose', methods=['POST'])
@_login_required
def packaging_compose():
    """Build the order email + a structured preview of the Word document.

    Returns the email envelope (to/cc/bcc/subject/body/attachment) plus the exact
    contents of the .docx (meta + line items + total), so the page can show a
    "Review & Send" panel that mirrors the attached Word document before sending.
    """
    try:
        sid = int(request.form.get('supplier_id', 0))
    except ValueError:
        return jsonify({'error': 'invalid supplier'}), 400
    supplier = _supplier(sid)
    if not supplier:
        return jsonify({'error': 'supplier not found'}), 404

    delivery_date = _resolve_delivery_date(supplier, request.form.get('delivery_date', ''))
    extra_note    = request.form.get('extra_note', '').strip()
    composed_by   = request.form.get('composed_by', session.get('role', '')).strip()

    chosen = _selected_order_items(sid)

    if not chosen:
        return jsonify({'error': 'No items selected. Set a quantity > 0 on at least one row.'}), 400

    try:
        subject, body = _compose_order(supplier, chosen, delivery_date, extra_note)
    except Exception as e:
        return jsonify({'error': f'Could not compose the order: {type(e).__name__}: {e}'}), 500
    cc = supplier.get('cc_emails', '')
    to = supplier['email']

    items = []
    for it in chosen:
        qty = int(it.get('qty') or 0)
        unit_price = float(it.get('unit_price') or 0)
        items.append({
            'product_code': it.get('product_code') or '',
            'name':         it.get('name_vi') or it.get('name_en') or '',
            'qty':          qty,
            'unit':         it.get('unit') or '',
            'unit_measure': it.get('unit_measure') or '',
            'unit_price':   unit_price,
            'line_total':   qty * unit_price,
        })
    total_qty, total_money = _order_totals(chosen)
    bcc = OWNER_BCC_EMAIL if (OWNER_BCC_EMAIL and OWNER_BCC_EMAIL.lower() != (to or '').lower()) else ''

    return jsonify({
        'ok': True,
        'subject':       subject,
        'body':          body,
        'to':            to,
        'cc':            cc,
        'bcc':           bcc,
        'docx_filename': _order_filename(supplier, delivery_date),
        # Word document content, mirroring _build_order_docx
        'brand':         _brand_name(supplier),
        'supplier_name': supplier.get('name') or '',
        'deliver_to':    _delivery_line(supplier),
        'contact':       (supplier.get('cafe_contacts') or '').strip(),
        'prepared_by':   composed_by or '-',
        'delivery_pretty': _format_delivery_date(delivery_date),
        'order_date':    datetime.now().strftime('%d/%m/%Y'),
        'note':          extra_note.strip(),
        'items':         items,
        'item_count':    len(chosen),
        'total_qty':     total_qty,
        'total_money':   total_money,
    })


@packaging_bp.route('/docx', methods=['POST'])
@_login_required
def packaging_docx():
    """Download the same Word order document used for supplier emails."""
    try:
        sid = int(request.form.get('supplier_id', 0))
    except ValueError:
        return jsonify({'error': 'invalid supplier'}), 400
    supplier = _supplier(sid)
    if not supplier:
        return jsonify({'error': 'supplier not found'}), 404

    delivery_date = _resolve_delivery_date(supplier, request.form.get('delivery_date', ''))
    extra_note    = request.form.get('extra_note', '').strip()
    composed_by   = request.form.get('composed_by', session.get('role','')).strip()
    chosen = _selected_order_items(sid)
    if not chosen:
        return jsonify({'error': 'No items selected.'}), 400

    docx_bytes = _build_order_docx(supplier, chosen, delivery_date, extra_note, composed_by)
    return send_file(
        io.BytesIO(docx_bytes),
        mimetype=DOCX_MIME,
        as_attachment=True,
        download_name=_order_filename(supplier, delivery_date),
    )


@packaging_bp.route('/send', methods=['POST'])
@_admin_required
def packaging_send():
    """Send order via Brevo HTTP API using the existing email_service config."""
    try:
        sid = int(request.form.get('supplier_id', 0))
    except ValueError:
        return jsonify({'error': 'invalid supplier'}), 400
    supplier = _supplier(sid)
    if not supplier:
        return jsonify({'error': 'supplier not found'}), 404

    delivery_date = _resolve_delivery_date(supplier, request.form.get('delivery_date', ''))
    extra_note    = request.form.get('extra_note', '').strip()
    composed_by   = request.form.get('composed_by', session.get('role','')).strip()

    chosen = _selected_order_items(sid)

    if not chosen:
        return jsonify({'error': 'No items selected.'}), 400

    # Build the email + Word document. Any failure here must return JSON — an
    # unhandled 500 renders an HTML page, which iOS Safari's resp.json() reports
    # as the cryptic "The string did not match the expected pattern."
    try:
        subject, body = _compose_order(supplier, chosen, delivery_date, extra_note)
        docx_bytes = _build_order_docx(supplier, chosen, delivery_date, extra_note, composed_by)
        filename = _order_filename(supplier, delivery_date)
    except ModuleNotFoundError as e:
        return jsonify({'error':
            'The server is missing the Word library (python-docx). On PythonAnywhere '
            'open a Bash console and run:  pip3 install --user python-docx  then Reload '
            f'the web app. [{e}]'}), 500
    except Exception as e:
        return jsonify({'error': f'Could not build the order document: {type(e).__name__}: {e}'}), 500

    # Send via Brevo
    try:
        import email_service
    except Exception:
        return jsonify({'error': 'email_service not available'}), 500

    settings = email_service.get_settings()
    if not email_service._is_configured(settings):
        return jsonify({'error':
            'Brevo not configured. Fill in the API key + sender email under Email Notifications first.'}), 400

    to_email = (supplier.get('email') or '').strip()
    if not to_email:
        return jsonify({'error': 'Supplier has no email address. Edit supplier info first.'}), 400

    total_qty, total_money = _order_totals(chosen)

    # HTML version stays short; the structured order is the attached DOCX.
    from html import escape as _esc
    html_body = (
        '<div style="font-family:Arial,sans-serif;font-size:14px;line-height:1.5;color:#222">'
        + f'<p>{_esc(body).replace(chr(10), "<br>")}</p>'
        + f'<p><strong>Attachment:</strong> {_esc(filename)}</p>'
        + '</div>'
    )

    cc_list = [c.strip() for c in (supplier.get('cc_emails') or '').split(',') if c.strip()]
    payload = {
        'sender':      {'name': settings.get('from_name') or _brand_name(supplier),
                         'email': settings['sender_email']},
        'to':          [{'email': to_email}],
        'subject':     subject,
        'htmlContent': html_body,
        'textContent': body,
        'attachment':   [{
            'content': base64.b64encode(docx_bytes).decode('ascii'),
            'name': filename,
        }],
    }
    if cc_list:
        payload['cc'] = [{'email': e} for e in cc_list]
    # Always BCC the owner so every order is archived without the supplier seeing it.
    if OWNER_BCC_EMAIL and OWNER_BCC_EMAIL.lower() != to_email.lower():
        payload['bcc'] = [{'email': OWNER_BCC_EMAIL}]

    ok, msg = email_service._brevo_post(payload, settings['brevo_api_key'])

    with _conn() as c:
        c.execute('''INSERT INTO packaging_orders
            (supplier_id, delivery_date, composed_by, send_channel, subject, body, payload_json, store_id)
            VALUES (?,?,?,?,?,?,?,?)''',
            (sid, delivery_date, composed_by,
             'brevo' if ok else f'brevo_failed',
             subject, body, json.dumps({
                 'items': chosen,
                 'docx_filename': filename,
                 'total_qty': total_qty,
                 'total_money': total_money,
             }), current_store_id()))

    if not ok:
        # Friendlier hint when Brevo bounces
        return jsonify({'error': f'Brevo send failed: {msg}'}), 500
    return jsonify({'ok': True, 'message': f'Word order emailed to {to_email}. Quantities reset to 0.'})


@packaging_bp.route('/log-action', methods=['POST'])
@_login_required
def packaging_log_action():
    """Record that the admin clicked the Gmail or mailto link so the order
    history shows it. No actual mail send happens here — Gmail/mailto opens
    in a new tab and the user clicks Send manually."""
    try:
        sid = int(request.form.get('supplier_id', 0))
    except ValueError:
        return jsonify({'error': 'invalid'}), 400
    channel = (request.form.get('channel', '') or '').strip().lower()
    if channel not in ('gmail', 'mailto', 'preview'):
        channel = 'preview'
    supplier = _supplier(sid)
    if not supplier:
        return jsonify({'error': 'supplier not found'}), 404
    subject = request.form.get('subject', '').strip()
    body    = request.form.get('body', '').strip()
    delivery_date = _resolve_delivery_date(supplier, request.form.get('delivery_date', ''))
    composed_by = request.form.get('composed_by', session.get('role','')).strip()
    with _conn() as c:
        c.execute('''INSERT INTO packaging_orders
            (supplier_id, delivery_date, composed_by, send_channel, subject, body, store_id)
            VALUES (?,?,?,?,?,?,?)''',
            (sid, delivery_date, composed_by, channel, subject, body, current_store_id()))
    return jsonify({'ok': True})


# ── Admin CRUD ───────────────────────────────────────────────────────────────

@packaging_bp.route('/supplier/<int:sid>/edit', methods=['POST'])
@_admin_required
def supplier_edit(sid):
    name = request.form.get('name', '').strip()
    if not name:
        return jsonify({'error': 'Supplier name required'}), 400
    delivery_days = request.form.get('delivery_days', '').strip().upper()
    if current_store_id() == SUBIACO_STORE_ID:
        delivery_days = SUBIACO_PACKAGING_DELIVERY_DAYS
    with _conn() as c:
        c.execute('''UPDATE packaging_suppliers SET
            name=?, email=?, phone=?, cc_emails=?, delivery_days=?,
            cafe_name=?, cafe_address=?, cafe_contacts=?, notes=?
            WHERE id=? AND store_id=?''',
            (name,
             request.form.get('email', '').strip(),
             request.form.get('phone', '').strip(),
             request.form.get('cc_emails', '').strip(),
             delivery_days,
             request.form.get('cafe_name', '').strip(),
             request.form.get('cafe_address', '').strip(),
             request.form.get('cafe_contacts', '').strip(),
             request.form.get('notes', '').strip(),
             sid, current_store_id()))
    return jsonify({'ok': True})


@packaging_bp.route('/supplier/add', methods=['POST'])
@_admin_required
def supplier_add():
    name = request.form.get('name', '').strip()
    if not name:
        return redirect(url_for('packaging.packaging_home'))
    delivery_days = request.form.get('delivery_days', '').strip().upper()
    if current_store_id() == SUBIACO_STORE_ID:
        delivery_days = SUBIACO_PACKAGING_DELIVERY_DAYS
    with _conn() as c:
        n = c.execute(
            'SELECT COALESCE(MAX(sort_order),-1)+1 AS n FROM packaging_suppliers WHERE store_id=?',
            (current_store_id(),)).fetchone()['n']
        cur = c.execute('''INSERT INTO packaging_suppliers
            (name, email, phone, cc_emails, delivery_days, cafe_name, cafe_contacts, sort_order, store_id)
            VALUES (?,?,?,?,?,?,?,?,?)''',
            (name,
             request.form.get('email', '').strip(),
             request.form.get('phone', '').strip(),
             request.form.get('cc_emails', '').strip(),
             delivery_days,
             request.form.get('cafe_name', '').strip() or 'MCQ Vietnamese Street Food — MIRRABOOKA',
             request.form.get('cafe_contacts', '').strip(),
             n, current_store_id()))
        new_id = cur.lastrowid
    return redirect(url_for('packaging.packaging_home', supplier=new_id))


@packaging_bp.route('/supplier/<int:sid>/delete', methods=['POST'])
@_admin_required
def supplier_delete(sid):
    with _conn() as c:
        c.execute('DELETE FROM packaging_suppliers WHERE id=? AND store_id=?',
                  (sid, current_store_id()))
    return redirect(url_for('packaging.packaging_home'))


@packaging_bp.route('/item/add', methods=['POST'])
@_admin_required
def item_add():
    try:
        sid = int(request.form.get('supplier_id', 0))
    except ValueError:
        return jsonify({'error': 'invalid supplier'}), 400
    name_en = request.form.get('name_en', '').strip()
    if not name_en:
        return jsonify({'error': 'Item name required'}), 400
    with _conn() as c:
        if not c.execute('SELECT 1 FROM packaging_suppliers WHERE id=? AND store_id=?',
                         (sid, current_store_id())).fetchone():
            return jsonify({'error': 'supplier not found'}), 404
        n = c.execute(
            'SELECT COALESCE(MAX(sort_order),-1)+1 AS n FROM packaging_items WHERE supplier_id=?',
            (sid,)).fetchone()['n']
        try:
            qty = max(0, int(request.form.get('default_qty', '0') or 0))
        except ValueError:
            qty = 0
        cur = c.execute('''INSERT INTO packaging_items
            (supplier_id, product_code, name_en, name_vi, unit, unit_measure,
             unit_price, default_qty, sort_order)
            VALUES (?,?,?,?,?,?,?,?,?)''',
            (sid,
             request.form.get('product_code', '').strip(),
             name_en,
             request.form.get('name_vi', '').strip(),
             request.form.get('unit', 'carton').strip() or 'carton',
             request.form.get('unit_measure', '').strip(),
             _float_form(request.form.get('unit_price'), 0.0),
             qty, n))
        new_id = cur.lastrowid
    return jsonify({'ok': True, 'id': new_id})


@packaging_bp.route('/item/<int:iid>/edit', methods=['POST'])
@_admin_required
def item_edit(iid):
    name_en = request.form.get('name_en', '').strip()
    if not name_en:
        return jsonify({'error': 'Item name required'}), 400
    try:
        qty = max(0, int(request.form.get('default_qty', '0') or 0))
    except ValueError:
        qty = 0
    with _conn() as c:
        if not c.execute('''SELECT 1 FROM packaging_items i
                            JOIN packaging_suppliers s ON s.id=i.supplier_id
                            WHERE i.id=? AND s.store_id=?''',
                         (iid, current_store_id())).fetchone():
            return jsonify({'error': 'item not found'}), 404
        c.execute('''UPDATE packaging_items SET
            product_code=?, name_en=?, name_vi=?, unit=?, unit_measure=?,
            unit_price=?, default_qty=?
            WHERE id=?''',
            (request.form.get('product_code', '').strip(),
             name_en,
             request.form.get('name_vi', '').strip(),
             request.form.get('unit', 'carton').strip() or 'carton',
             request.form.get('unit_measure', '').strip(),
             _float_form(request.form.get('unit_price'), 0.0),
             qty, iid))
    return jsonify({'ok': True})


@packaging_bp.route('/item/<int:iid>/delete', methods=['POST'])
@_admin_required
def item_delete(iid):
    with _conn() as c:
        c.execute('''DELETE FROM packaging_items
            WHERE id=? AND supplier_id IN (
              SELECT id FROM packaging_suppliers WHERE store_id=?
            )''', (iid, current_store_id()))
    return jsonify({'ok': True})


@packaging_bp.route('/history')
@_login_required
def packaging_history():
    sid = request.args.get('supplier', '')
    scope, sp = store_filter_clause('o')
    with _conn() as c:
        q = ('SELECT o.*, s.name AS supplier_name FROM packaging_orders o '
             'LEFT JOIN packaging_suppliers s ON s.id=o.supplier_id '
             f'WHERE {scope}')
        params = list(sp)
        if sid:
            q += ' AND o.supplier_id=?'; params.append(int(sid))
        q += ' ORDER BY o.id DESC LIMIT 50'
        rows = [dict(r) for r in c.execute(q, params).fetchall()]
    return render_template('packaging_history.html', orders=rows)
